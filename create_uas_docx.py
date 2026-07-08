import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_table_row(table, data):
    row_cells = table.add_row().cells
    for i, item in enumerate(data):
        row_cells[i].text = str(item)

doc = docx.Document()

# Title
title = doc.add_heading('TEMPLATE PENGUMPULAN UAS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle1 = doc.add_paragraph('Mata Kuliah Pemrograman Mobile')
subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2 = doc.add_paragraph('Proyek Individu Aplikasi Mobile Sederhana')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dokumen ini digunakan sebagai format laporan pengumpulan Ujian Akhir Semester. Isian laporan dibuat singkat, jelas, dan mudah diperiksa. Setiap bagian perlu diisi sesuai kondisi aplikasi yang dikembangkan.')

# A. Identitas Mahasiswa dan Proyek
doc.add_heading('A. Identitas Mahasiswa dan Proyek', level=1)
doc.add_paragraph('Bagian ini berisi identitas pembuat aplikasi dan informasi dasar proyek. Data ditulis lengkap agar proses pemeriksaan lebih mudah dilakukan.')

table_a = doc.add_table(rows=1, cols=2)
table_a.style = 'Table Grid'
hdr_cells = table_a.rows[0].cells
hdr_cells[0].text = 'Keterangan'
hdr_cells[1].text = 'Isian'

identitas = [
    ('Nama Mahasiswa', 'Dimas Ulung Septiaji'),
    ('NIM', '240103014'),
    ('Program Studi', 'Teknik Informatika'),
    ('Kelas', '-'),
    ('Dosen Pengampu', 'Dr. Nurchim, S. Kom., M. Kom'),
    ('Mata Kuliah', 'Pemrograman Mobile'),
    ('Judul Aplikasi', 'Dims Movie'),
    ('Platform API Eksternal', 'TMDB (The Movie Database) API & Custom Backend PHP')
]
for item in identitas:
    add_table_row(table_a, item)

# B. Deskripsi Aplikasi Mobile Sederhana
doc.add_heading('B. Deskripsi Aplikasi Mobile Sederhana', level=1)
doc.add_paragraph('1. Nama Aplikasi')
doc.add_paragraph('Dims Movie')

doc.add_paragraph('2. Latar Belakang Singkat')
doc.add_paragraph('Aplikasi ini dibuat untuk memberikan kemudahan bagi penggemar film dalam mencari informasi film populer saat ini serta mengelola daftar tontonan pribadi. Dengan mengintegrasikan API pihak ketiga (TMDB), aplikasi dapat menampilkan data film secara real-time yang informatif.')

doc.add_paragraph('3. Tujuan Aplikasi')
doc.add_paragraph('Menyediakan platform mobile untuk katalog film interaktif yang memungkinkan pengguna mencari film, melihat detailnya, dan mencatat film-film yang ingin ditonton beserta catatan pribadinya dalam satu aplikasi yang tersinkronisasi ke server.')

doc.add_paragraph('4. Pengguna yang Dituju')
doc.add_paragraph('Masyarakat umum, penggemar film, dan individu yang hobi menonton serta ingin membuat riwayat/daftar tontonannya.')

doc.add_paragraph('5. Fungsi Utama Aplikasi')
table_b = doc.add_table(rows=1, cols=3)
table_b.style = 'Table Grid'
hdr_cells = table_b.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Fungsi Utama'
hdr_cells[2].text = 'Keterangan Singkat'
fungsi_utama = [
    ('1', 'Menampilkan film populer', 'Mengambil dan menampilkan daftar film terpopuler dari TMDB API'),
    ('2', 'Mencari film', 'Mencari film berdasarkan judul secara real-time'),
    ('3', 'Melihat detail film', 'Menampilkan poster, rating, dan deskripsi film yang dipilih'),
    ('4', 'Mengelola Watchlist', 'Menambahkan, mengubah catatan, dan menghapus film dari daftar tontonan pribadi'),
    ('5', '', '')
]
for item in fungsi_utama:
    add_table_row(table_b, item)

# C. Platform API Eksternal
doc.add_heading('C. Platform API Eksternal', level=1)
table_c1 = doc.add_table(rows=1, cols=2)
table_c1.style = 'Table Grid'
hdr_cells = table_c1.rows[0].cells
hdr_cells[0].text = 'Keterangan'
hdr_cells[1].text = 'Isian'
api_eksternal = [
    ('Nama Platform API', 'The Movie Database (TMDB) API'),
    ('URL Dokumentasi Resmi', 'https://developer.themoviedb.org/docs'),
    ('URL API yang Digunakan', 'https://api.themoviedb.org/3/movie/popular, https://api.themoviedb.org/3/search/movie'),
    ('Jenis Data yang Diambil', 'JSON'),
    ('API Key', '84ec6a2842323c09bdd36a99ecb21a2c')
]
for item in api_eksternal:
    add_table_row(table_c1, item)

doc.add_paragraph('\nData yang Diambil dari Platform API')
table_c2 = doc.add_table(rows=1, cols=4)
table_c2.style = 'Table Grid'
hdr_cells = table_c2.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Data'
hdr_cells[2].text = 'Contoh Isi Data'
hdr_cells[3].text = 'Keterangan'
data_diambil = [
    ('1', 'title', 'Inception', 'Judul film'),
    ('2', 'poster_path', '/oYuLEt3zVCKq57qu2F8dT7NIa6f.jpg', 'Path gambar poster'),
    ('3', 'overview', 'A thief who steals corporate secrets...', 'Deskripsi/Sinopsis cerita'),
    ('4', 'vote_average', '8.8', 'Nilai rating dari TMDB'),
    ('5', 'id', '27205', 'ID film unik TMDB')
]
for item in data_diambil:
    add_table_row(table_c2, item)

doc.add_paragraph('\nCara Data Digunakan dalam Aplikasi')
doc.add_paragraph('Data dari TMDB API diambil dengan method GET menggunakan token Authorization. Setelah JSON diterima, data di-parsing menjadi list model Movie untuk ditampilkan pada list/grid. Saat pengguna memilih untuk menyimpan ke watchlist, field tersebut (id, title, poster_path, rating) dikirim ke backend database server sendiri untuk disimpan.')

# D. Database Server
doc.add_heading('D. Database Server', level=1)
table_d1 = doc.add_table(rows=1, cols=2)
table_d1.style = 'Table Grid'
hdr_cells = table_d1.rows[0].cells
hdr_cells[0].text = 'Keterangan'
hdr_cells[1].text = 'Isian'
db_info = [
    ('Nama Database', 'dmovie_db (Asumsi)'),
    ('Jenis Database', 'MySQL / MariaDB'),
    ('Alamat Server Database', 'dmovie.dimas-server.my.id'),
    ('Nama Tabel atau Koleksi Utama', 'watchlist'),
    ('Fungsi Tabel atau Koleksi Utama', 'Menyimpan daftar film yang dimasukkan ke watchlist oleh pengguna')
]
for item in db_info:
    add_table_row(table_d1, item)

doc.add_paragraph('\nStruktur Tabel')
table_d2 = doc.add_table(rows=1, cols=4)
table_d2.style = 'Table Grid'
hdr_cells = table_d2.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Tabel'
hdr_cells[2].text = 'Fungsi'
hdr_cells[3].text = 'Jumlah Field'
add_table_row(table_d2, ('1', 'watchlist', 'Menyimpan data film watchlist', '6'))
add_table_row(table_d2, ('2', '', '', ''))
add_table_row(table_d2, ('3', '', '', ''))

doc.add_paragraph('\nField Relasi Tabel')
table_d3 = doc.add_table(rows=1, cols=4)
table_d3.style = 'Table Grid'
hdr_cells = table_d3.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Field'
hdr_cells[2].text = 'Tipe Data'
hdr_cells[3].text = 'Keterangan'
fields = [
    ('1', 'id', 'INT', 'Primary Key'),
    ('2', 'movie_id', 'INT', 'ID film dari TMDB API'),
    ('3', 'title', 'VARCHAR', 'Judul film'),
    ('4', 'poster_path', 'VARCHAR', 'Path poster'),
    ('5', 'rating', 'DOUBLE', 'Nilai rating'),
    ('6', 'notes', 'TEXT', 'Catatan pribadi')
]
for item in fields:
    add_table_row(table_d3, item)

doc.add_paragraph('\nTangkapan Layar Database Server')
doc.add_paragraph('[Lampirkan tangkapan layar database server di sini]')
doc.add_paragraph('[Lampirkan tangkapan layar struktur tabel di sini]')
doc.add_paragraph('[Lampirkan tangkapan layar data tersimpan di sini]')

# E. Daftar URL REST API
doc.add_heading('E. Daftar URL REST API', level=1)
table_e1 = doc.add_table(rows=1, cols=2)
table_e1.style = 'Table Grid'
hdr_cells = table_e1.rows[0].cells
hdr_cells[0].text = 'Keterangan'
hdr_cells[1].text = 'Isian'
api_rest = [
    ('Base URL API', 'https://dmovie.dimas-server.my.id'),
    ('Format Respons', 'JSON'),
    ('Status Akses', 'Publik')
]
for item in api_rest:
    add_table_row(table_e1, item)

doc.add_paragraph('\nEndpoint REST API')
table_e2 = doc.add_table(rows=1, cols=4)
table_e2.style = 'Table Grid'
hdr_cells = table_e2.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Kebutuhan'
hdr_cells[2].text = 'Method'
hdr_cells[3].text = 'URL Endpoint'
endpoints = [
    ('1', 'Menampilkan seluruh data', 'GET', '/watchlist.php'),
    ('2', 'Menampilkan detail data', 'GET', 'TMDB API (Search/Popular)'),
    ('3', 'Menambahkan data', 'POST', '/watchlist.php'),
    ('4', 'Mengubah data', 'PUT', '/watchlist.php'),
    ('5', 'Menghapus data', 'DELETE', '/watchlist.php?movie_id={id}')
]
for item in endpoints:
    add_table_row(table_e2, item)

doc.add_paragraph('\nContoh Respons REST API')
table_e3 = doc.add_table(rows=1, cols=2)
table_e3.style = 'Table Grid'
hdr_cells = table_e3.rows[0].cells
hdr_cells[0].text = 'Proses'
hdr_cells[1].text = 'Contoh Respons'
respons = [
    ('SELECT', '{ "data": [ { "movie_id": 27205, "title": "Inception" } ] }'),
    ('INSERT', '{ "status": true, "message": "Berhasil menambahkan ke watchlist" }'),
    ('UPDATE', '{ "status": true, "message": "Berhasil mengupdate watchlist" }'),
    ('DELETE', '{ "status": true, "message": "Berhasil menghapus dari watchlist" }')
]
for item in respons:
    add_table_row(table_e3, item)

# F. Tautan Aplikasi
doc.add_heading('F. Tautan Aplikasi, APK, dan Kode Sumber', level=1)
table_f = doc.add_table(rows=1, cols=2)
table_f.style = 'Table Grid'
hdr_cells = table_f.rows[0].cells
hdr_cells[0].text = 'Keterangan'
hdr_cells[1].text = 'Tautan atau Isian'
tautan = [
    ('URL Repositori Aplikasi Mobile', '-'),
    ('URL Repositori Server', '-'),
    ('URL File APK Jika Berbasis Android', '-'),
    ('URL Aplikasi Jika Berbasis Web Mobile', '-'),
    ('URL Server Aktif', 'https://dmovie.dimas-server.my.id/watchlist.php'),
    ('Akun Uji, Jika Ada', '-'),
    ('Kata Sandi Uji, Jika Ada', '-'),
    ('Catatan Instalasi', 'Jalankan menggunakan `flutter run`')
]
for item in tautan:
    add_table_row(table_f, item)

# G. Dokumentasi Alur Kerja Aplikasi
doc.add_heading('G. Dokumentasi Alur Kerja Aplikasi', level=1)
doc.add_paragraph('1. Ringkasan Alur Kerja')
table_g = doc.add_table(rows=1, cols=3)
table_g.style = 'Table Grid'
hdr_cells = table_g.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Langkah Kerja'
hdr_cells[2].text = 'Penjelasan Singkat'
alur = [
    ('1', 'Buka Aplikasi', 'Aplikasi mengambil data film populer dari TMDB API.'),
    ('2', 'Cari Film', 'Pengguna mencari film di Search Screen.'),
    ('3', 'Lihat Detail', 'Pengguna menekan film untuk melihat detail.'),
    ('4', 'Tambah Watchlist', 'Pengguna menekan simpan, data POST ke backend.'),
    ('5', 'Buka Watchlist', 'Aplikasi memanggil GET ke backend untuk menampilkan data.'),
    ('6', 'Ubah/Hapus Data', 'Pengguna bisa update/hapus data (PUT/DELETE) disinkronisasi ke server.'),
    ('7', '', ''),
    ('8', '', '')
]
for item in alur:
    add_table_row(table_g, item)

doc.add_paragraph('\n2. Alur Pengambilan Data dari Platform API')
doc.add_paragraph('Aplikasi mengirimkan HTTP GET request dengan header token Authorization ke endpoint TMDB. Response JSON di-decode menjadi model Movie dan ditampilkan melalui state management.')

doc.add_paragraph('\n3. Alur Penyimpanan Data ke Database Server')
doc.add_paragraph('Saat pengguna menekan tombol tambah watchlist, aplikasi membuat JSON data film dan mengirim via HTTP POST ke backend. Server PHP memproses dan INSERT ke database.')

doc.add_paragraph('\n4. Alur Menampilkan Data dari Database Server')
doc.add_paragraph('Aplikasi mengirim HTTP GET ke endpoint backend. Server merespons hasil SELECT tabel watchlist dalam format JSON array yang kemudian dirender oleh aplikasi.')

doc.add_paragraph('\n5. Alur Mengubah Data')
doc.add_paragraph('Saat pengguna mengedit notes, aplikasi mengirimkan HTTP PUT berisi data pembaruan. Backend memproses UPDATE berdasarkan movie_id.')

doc.add_paragraph('\n6. Alur Menghapus Data')
doc.add_paragraph('Pengguna menekan hapus, aplikasi mengirimkan HTTP DELETE ke backend PHP dengan parameter movie_id. Backend menjalankan query DELETE.')

# H. Tangkapan Layar Aplikasi
doc.add_heading('H. Tangkapan Layar Aplikasi', level=1)
table_h = doc.add_table(rows=1, cols=4)
table_h.style = 'Table Grid'
hdr_cells = table_h.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Halaman atau Proses'
hdr_cells[2].text = 'Keterangan'
hdr_cells[3].text = 'Tangkapan Layar'
ss = [
    ('1', 'Halaman utama', 'Daftar film populer', '[Lampirkan SS]'),
    ('2', 'Pengambilan data dari platform API', 'Search results dari TMDB', '[Lampirkan SS]'),
    ('3', 'Daftar data tersimpan', 'Watchlist page', '[Lampirkan SS]'),
    ('4', 'Detail data', 'Movie detail page', '[Lampirkan SS]'),
    ('5', 'Proses simpan data', 'Tambah ke watchlist', '[Lampirkan SS]'),
    ('6', 'Proses ubah data', 'Edit notes', '[Lampirkan SS]'),
    ('7', 'Proses hapus data', 'Hapus dari watchlist', '[Lampirkan SS]'),
    ('8', 'Pesan berhasil atau gagal', 'Snackbar feedback', '[Lampirkan SS]')
]
for item in ss:
    add_table_row(table_h, item)

# I. Hasil Pengujian Fitur
doc.add_heading('I. Hasil Pengujian Fitur', level=1)
table_i = doc.add_table(rows=1, cols=4)
table_i.style = 'Table Grid'
hdr_cells = table_i.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Fitur yang Diuji'
hdr_cells[2].text = 'Status'
hdr_cells[3].text = 'Keterangan'
fitur = [
    ('1', 'Mengambil data dari platform API eksternal', 'Berhasil', 'Berjalan lancar'),
    ('2', 'Menampilkan data dari platform API eksternal', 'Berhasil', 'UI render sempurna'),
    ('3', 'Menyimpan data ke database server', 'Berhasil', 'Data masuk ke PHP/MySQL'),
    ('4', 'Menampilkan data dari database server', 'Berhasil', 'List ter-update'),
    ('5', 'Menambahkan data melalui REST API', 'Berhasil', 'POST berhasil'),
    ('6', 'Mengubah data melalui REST API', 'Berhasil', 'PUT berhasil'),
    ('7', 'Menghapus data melalui REST API', 'Berhasil', 'DELETE berhasil'),
    ('8', 'Navigasi antarmuka aplikasi', 'Berhasil', 'Transisi layar lancar')
]
for item in fitur:
    add_table_row(table_i, item)

# J. Catatan Tambahan
doc.add_heading('J. Catatan Tambahan', level=1)
doc.add_paragraph('Aplikasi membutuhkan koneksi internet yang stabil untuk mengakses API TMDB dan custom backend server. Harap pastikan server dmovie.dimas-server.my.id sedang menyala pada saat penilaian.')

doc.save('Laporan_UAS_Dims_Movie.docx')
